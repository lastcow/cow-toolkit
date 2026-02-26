"""Attachment downloader and content extractor.

Supports:
  Text:   .docx, .pdf, .txt, .md, .py, .html, .csv, and other plain text
  Images: .jpg/.jpeg, .png, .gif, .webp, .bmp, .tiff (standalone attachments)
  Mixed:  DOCX/PDF with embedded images — images are extracted and analyzed
"""

import io
import os
import re
import subprocess
import tempfile
from pathlib import Path
from urllib.parse import unquote_plus


def _get_token() -> str:
    return os.environ.get("CANVAS_API_TOKEN", "")


# ── Image analysis via claude CLI ────────────────────────────────────────────

def _analyze_image_file(path: str) -> str:
    """Use claude CLI to OCR / describe an image file.
    Returns extracted text, or a fallback description on failure.
    """
    try:
        result = subprocess.run(
            [
                "claude", "--print",
                "--dangerously-skip-permissions",
                "--add-dir", str(Path(path).parent),
                "--model", "claude-haiku-4-5",   # fast model for image OCR
                (
                    f"Analyze the image file at {path}. "
                    "Extract ALL readable text exactly as written. "
                    "If there are charts/figures, briefly describe them. "
                    "Keep response concise. No preamble."
                ),
            ],
            capture_output=True, text=True, timeout=45,
        )
        text = result.stdout.strip()
        if text and len(text) > 5:
            return text
        return "[Image: no readable text detected]"
    except subprocess.TimeoutExpired:
        return "[Image: analysis timed out]"
    except FileNotFoundError:
        return "[Image: claude CLI not available for OCR]"
    except Exception as e:
        return f"[Image: analysis failed — {e}]"


def _analyze_image_bytes(data: bytes, suffix: str = ".png") -> str:
    """Save image bytes to temp file, analyze, and clean up."""
    with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tf:
        tf.write(data)
        tmp_path = tf.name
    try:
        return _analyze_image_file(tmp_path)
    finally:
        try:
            os.unlink(tmp_path)
        except Exception:
            pass


def _image_metadata(data: bytes) -> str:
    """Return basic PIL metadata if vision analysis unavailable."""
    try:
        from PIL import Image
        img = Image.open(io.BytesIO(data))
        return f"[Image {img.width}×{img.height} {img.mode}]"
    except Exception:
        return "[Image: preview unavailable]"


# ── Main entry point ──────────────────────────────────────────────────────────

def fetch_attachment_content(att) -> dict:
    """Download and extract full content from a Canvas attachment object.

    Returns dict:
        filename       str
        content_type   str
        size           int
        text           str    (all extracted text, including image OCR)
        error          str    (non-empty on failure)
    """
    import requests

    filename     = unquote_plus(getattr(att, "filename", "unknown"))
    content_type = getattr(att, "content-type", "") or ""
    size         = getattr(att, "size", 0) or 0
    url          = getattr(att, "url", "")

    result = {
        "filename":     filename,
        "content_type": content_type,
        "size":         size,
        "text":         "",
        "error":        "",
    }

    if not url:
        result["error"] = "No download URL available"
        return result

    if size and size > 10_000_000:
        result["error"] = f"File too large ({size // 1_000_000} MB) to preview"
        return result

    try:
        token = _get_token()
        resp  = requests.get(
            url,
            headers={"Authorization": f"Bearer {token}"},
            timeout=30,
        )
        resp.raise_for_status()
        data = resp.content
    except Exception as e:
        result["error"] = f"Download failed: {e}"
        return result

    lname = filename.lower()
    ct    = content_type.lower()

    try:
        # ── DOCX ──────────────────────────────────────────────────────────
        if "wordprocessingml" in ct or lname.endswith(".docx"):
            result["text"] = _extract_docx(data)

        # ── PDF ───────────────────────────────────────────────────────────
        elif "pdf" in ct or lname.endswith(".pdf"):
            result["text"] = _extract_pdf(data)

        # ── Standalone images ─────────────────────────────────────────────
        elif _is_image(ct, lname):
            suffix = Path(lname).suffix or ".jpg"
            result["text"] = _analyze_image_bytes(data, suffix)

        # ── Plain text ────────────────────────────────────────────────────
        elif _is_plaintext(ct, lname):
            result["text"] = data.decode("utf-8", errors="replace")

        else:
            result["error"] = f"Unsupported format: {filename}"

    except Exception as e:
        result["error"] = f"Parse error: {e}"

    return result


# ── Helpers ───────────────────────────────────────────────────────────────────

_IMAGE_EXTS = {
    ".jpg", ".jpeg", ".png", ".gif", ".webp",
    ".bmp", ".tiff", ".tif", ".avif", ".heic",
}
_TEXT_EXTS = {
    ".txt", ".md", ".py", ".js", ".ts", ".java", ".c", ".cpp", ".h",
    ".css", ".html", ".htm", ".csv", ".json", ".xml", ".sh", ".yaml",
    ".yml", ".rst", ".rb", ".go", ".rs", ".kt", ".swift",
}


def _is_image(ct: str, fname: str) -> bool:
    return ct.startswith("image/") or Path(fname).suffix.lower() in _IMAGE_EXTS


def _is_plaintext(ct: str, fname: str) -> bool:
    return "text/" in ct or Path(fname).suffix.lower() in _TEXT_EXTS


def _extract_docx(data: bytes) -> str:
    import docx as docxlib

    doc    = docxlib.Document(io.BytesIO(data))
    parts  = []

    # ── Text paragraphs ───────────────────────────────────────────────────
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            parts.append(t)

    # ── Tables ────────────────────────────────────────────────────────────
    for tbl in doc.tables:
        rows = []
        for row in tbl.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                rows.append(" | ".join(cells))
        if rows:
            parts.append("\n── Table ──\n" + "\n".join(rows))

    # ── Embedded images ───────────────────────────────────────────────────
    img_count = 0
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            img_count += 1
            try:
                img_data  = rel.target_part.blob
                img_name  = rel.target_part.partname.split("/")[-1].lower()
                suffix    = Path(img_name).suffix or ".png"
                ocr_text  = _analyze_image_bytes(img_data, suffix)
                parts.append(f"\n── Embedded Image {img_count} ({img_name}) ──\n{ocr_text}")
            except Exception as e:
                parts.append(f"\n── Embedded Image {img_count} ──\n[extraction failed: {e}]")

    return "\n".join(parts)


def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader     = PdfReader(io.BytesIO(data))
    page_texts = []

    for page_num, page in enumerate(reader.pages, 1):
        text = (page.extract_text() or "").strip()

        # Try to extract embedded images from this page
        img_parts = []
        try:
            for img_obj in page.images:
                img_data = img_obj.data
                suffix   = Path(img_obj.name.lower()).suffix if img_obj.name else ".png"
                if not suffix:
                    suffix = ".png"
                ocr_text = _analyze_image_bytes(img_data, suffix)
                if ocr_text and "[Image:" not in ocr_text:
                    img_parts.append(f"[Image in page {page_num}]: {ocr_text}")
        except Exception:
            pass  # image extraction not always available

        if text:
            page_texts.append(text)
        if img_parts:
            page_texts.extend(img_parts)

    return "\n\n".join(page_texts)


def format_size(size: int) -> str:
    if size < 1024:
        return f"{size} B"
    elif size < 1_048_576:
        return f"{size / 1024:.1f} KB"
    else:
        return f"{size / 1_048_576:.1f} MB"
